"""
Regenerate QA DOCX files for AGENCY15-316 with AI-reviewed findings pre-filled.
Overwrites the blank files created by create-qa-docs.py.

Usage:
  python scripts/populate-qa-agency15-316.py
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TICKET  = 'AGENCY15-316'
QA_ROOT = str(Path(__file__).parent.parent / 'qa-reports')
OUT_DIR = os.path.join(QA_ROOT, TICKET)

# --- Colors ---
NAVY    = '1F4E79'
BLUE    = '2E86AB'
LGREY   = 'D5D8DC'
ALTROW  = 'EBF5FB'
WHITE   = 'FFFFFF'
SEV_HIGH   = 'FDECEA'
SEV_MED    = 'FEF9E7'
SEV_LOW    = 'EAF2FF'
SEV_INFO   = 'E8F8F5'
NAVY_C  = RGBColor(0x1F, 0x4E, 0x79)
BLUE_C  = RGBColor(0x2E, 0x86, 0xAB)
WHITE_C = RGBColor(0xFF, 0xFF, 0xFF)
DARK_C  = RGBColor(0x1A, 0x1A, 0x1A)
GREY_C  = RGBColor(0x66, 0x66, 0x66)
LGREY_C = RGBColor(0x88, 0x88, 0x88)
RED_C   = RGBColor(0xC0, 0x39, 0x2B)
ORG_C   = RGBColor(0xD3, 0x5A, 0x00)
BLU_C   = RGBColor(0x1A, 0x5C, 0xBF)

SEV_COLOR = {
    'HIGH':   ('C0392B', RED_C),
    'MEDIUM': ('D35A00', ORG_C),
    'LOW':    ('1A5CBF', BLU_C),
    'INFO':   ('27AE60', RGBColor(0x27, 0xAE, 0x60)),
}

# --- XML helpers ---
def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _cell_w(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:w'), str(dxa))
    tcW.set(qn('w:type'), 'dxa')

def _tbl_w(table, dxa):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(dxa))
    tblW.set(qn('w:type'), 'dxa')

def _row_h(row, val=380):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(val))
    h.set(qn('w:hRule'), 'atLeast')
    trPr.append(h)

def _para_border_bottom(para, color=NAVY):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

# --- Text helpers ---
def _t(doc, text, size=9.5, bold=False, italic=False, color=None,
       before=0, after=2, indent=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if align:  p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def title(doc, text):
    return _t(doc, text, size=20, bold=True, color=NAVY_C, before=0, after=4)

def subtitle(doc, text):
    return _t(doc, text, size=10, italic=True, color=GREY_C, before=0, after=10)

def section(doc, text):
    p = _t(doc, text, size=12, bold=True, color=NAVY_C, before=12, after=3)
    _para_border_bottom(p)
    return p

def subsection(doc, text):
    return _t(doc, text, size=10, bold=True, color=BLUE_C, before=6, after=2)

def check(doc, text):
    return _t(doc, '☐  ' + text, size=9.5, color=DARK_C, before=1, after=1, indent=0.2)

def note(doc, text):
    return _t(doc, text, size=8.5, italic=True, color=LGREY_C, before=2, after=2, indent=0.2)

def spacer(doc):
    _t(doc, '', size=8, before=0, after=4)

PAGE_W = 9072

def _hdr_cell(cell, text, w):
    _cell_w(cell, w); _cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE_C

def _data_cell(cell, text, w, bg=WHITE, bold=False, color=None):
    _cell_w(cell, w); _cell_bg(cell, bg)
    if text:
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(8.5)
        r.font.color.rgb = color if color else GREY_C
        if bold: r.bold = True

def info_table(doc, rows):
    W = [2000, PAGE_W - 2000]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (label, val) in enumerate(rows):
        r = tbl.rows[i]
        _cell_w(r.cells[0], W[0]); _cell_bg(r.cells[0], LGREY)
        _cell_w(r.cells[1], W[1]); _cell_bg(r.cells[1], WHITE)
        rr = r.cells[0].paragraphs[0].add_run(label)
        rr.bold = True; rr.font.size = Pt(9)
        if val:
            rr2 = r.cells[1].paragraphs[0].add_run(val)
            rr2.font.size = Pt(9)
        _row_h(r, 340)

def session_log(doc, n=3):
    W = [700, 2000, 2500, PAGE_W - 700 - 2000 - 2500]
    hdrs = ['Run', 'Date', 'Reviewed by', 'Notes']
    tbl = doc.add_table(rows=1 + n, cols=4)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, W)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 400)
    prefill = [('1', '2026-08-31', 'Claude AI (automated)', 'Initial AI QA — 13 findings')]
    for i in range(n):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        vals = prefill[i] if i < len(prefill) else ('', '', '', '')
        for j, (w, v) in enumerate(zip(W, vals)):
            _data_cell(row.cells[j], v, w, bg)
        _row_h(row, 420)

def findings_table_populated(doc, headers, col_widths, findings):
    """
    findings: list of dicts with keys:
      id, severity, run, cols (list matching headers length), status
    """
    all_hdrs  = ['ID', 'Sev', 'Run'] + headers + ['Status']
    sev_w     = 700
    run_w     = 500
    status_w  = 900
    id_w      = 700
    inner_total = PAGE_W - id_w - sev_w - run_w - status_w
    ratio       = inner_total / sum(col_widths)
    inner_ws    = [max(800, int(w * ratio)) for w in col_widths]
    inner_ws[-1] += inner_total - sum(inner_ws)
    all_ws = [id_w, sev_w, run_w] + inner_ws + [status_w]

    n = max(len(findings), 5)
    tbl = doc.add_table(rows=1 + n, cols=len(all_hdrs))
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)

    for i, (h, w) in enumerate(zip(all_hdrs, all_ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)

    for i in range(n):
        row = tbl.rows[i + 1]
        if i < len(findings):
            f   = findings[i]
            sev = f['severity']
            bg  = SEV_HIGH if sev == 'HIGH' else SEV_MED if sev == 'MEDIUM' else SEV_LOW if sev == 'LOW' else SEV_INFO
            _, sev_c = SEV_COLOR[sev]
            vals = [f['id'], sev, f.get('run', '1')] + f['cols'] + [f.get('status', 'Open')]
            colors = [None, sev_c, None] + [None]*len(f['cols']) + [None]
            bolds  = [True, True, False] + [False]*len(f['cols']) + [False]
        else:
            bg     = ALTROW if i % 2 == 1 else WHITE
            vals   = [''] * len(all_hdrs)
            colors = [None] * len(all_hdrs)
            bolds  = [False] * len(all_hdrs)

        for j, (w, v) in enumerate(zip(all_ws, vals)):
            _data_cell(row.cells[j], v, w, bg, bold=bolds[j], color=colors[j])
        _row_h(row, 520)


# =============================================================================
# Findings data
# =============================================================================

CONTENT_FINDINGS = [
    {
        'id': 'C-01', 'severity': 'HIGH', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:998-1003\n(both hero imgs)',
            'alt="特発性好酸球増加症候群診療の参照ガイド 令和7年度初版"\n(actual banner content per Figma Tab 1 desktop)',
            'alt="NATRON試験 二重盲検投与期（24週間）の成績"\n(copy-paste from different article — does not describe this page)',
        ],
        'status': 'Open',
    },
    {
        'id': 'C-02', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:1889-1898\n(Tab 2 mobile nav bar)',
            '3 TOC links: 診断・鑑別 (#h2-0), 治療戦略 (#h2-1), 管理・予後 (#h2-2)',
            '2 links only; labels wrong (診療の参照ガイドの概要 / クリニカルクエスチョン); 管理・予後 (#h2-2) missing entirely',
        ],
        'status': 'Open',
    },
]

CODING_FINDINGS = [
    {
        'id': 'K-01', 'severity': 'HIGH', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:64\n<title>',
            '<title></title>\n(empty — no page title set)',
            '<title>特発性好酸球増加症候群 診療の参照ガイド ｜ MediChannel</title>\n(45–65 chars, aligned with h1)',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-02', 'severity': 'HIGH', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:990 + :1884\n(two role="main")',
            '<div id="main" role="main"> (line 990)\nAND\n<div id="clinical-questions" role="main"> (line 1884)',
            'Keep role="main" on #main only.\nRemove role="main" from #clinical-questions.\nRule: exactly one main landmark per page.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-03', 'severity': 'HIGH', 'run': '1',
        'cols': [
            'fsn_hes_article03.html\n(no h1 anywhere in editable area)',
            'No <h1> element present in the page.',
            'Add <h1> for the hero/banner title "特発性好酸球増加症候群診療の参照ガイド".\nMay be visually hidden with clip-path if design does not show it.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-04', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:998 + :1003\n(desktop and mobile hero imgs)',
            'loading="lazy"\n(on both hero images above the fold)',
            'loading="eager" fetchpriority="high"\n(LCP/hero images must be eager per guidelines)',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-05', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:1482\n(inside .cst-off-label-note)',
            '<div class="cst-hero-wrapper">\n  <div class="cst-abbreviation-row">...\n(cst-hero-wrapper adds 20 px bottom padding — incorrect here)',
            'Replace with a neutral wrapper or remove the outer div.\nDo not reuse cst-hero-wrapper outside the hero section.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-06', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'page.css — anchor targets\n#gaiyou, #shindan, #chiryo, #monitoring\n(secondary sticky TOC links)',
            'No scroll-margin-top defined on any anchor-target element.\n(sticky TOC is top:0 desktop / top:64px mobile)',
            'Add:\n#gaiyou, #shindan, #chiryo, #monitoring {\n  scroll-margin-top: 64px;\n}\n@media (max-width:767px) { same: 128px }',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-07', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:2057\n(Tab 2 footer)',
            '<p class="cst-footer-note__link">電子添文を見る</p>\n(styled as link but not interactive — not an <a>)',
            '<a class="cst-footer-note__link" href="[PI URL]">電子添文を見る</a>\n(match Tab 1 conclusion button — must be a real anchor)',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-08', 'severity': 'MEDIUM', 'run': '1',
        'cols': [
            'html:1571 (Tab 1) vs html:1978 (Tab 2)\n(approval-notice element)',
            'Tab 1: <strong class="cst-approval-notice">本参照ガイド発行後の2026年5月に…</strong>\nTab 2: <h4 class="cst-update__title">同テキスト</h4>',
            'Use the same element type in both tabs.\nIf this is a heading, use <h4> in both.\nIf it is notice text, use <p> or <div> in both.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-09', 'severity': 'LOW', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:1029\n(.cst-toc-secondary)',
            '<div class="cst-toc-secondary">\n(no role, no accessible name)',
            'Add role="navigation" aria-label="目次" to .cst-toc-secondary\nso screen readers can identify the TOC landmark.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-10', 'severity': 'LOW', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:1576\n(PI box label in Tab 1)',
            '<strong class="cst-gene-name">4. 効能又は効果（抜粋）</strong>\n(cst-gene-name applied to a section label, not a gene name)',
            'Replace cst-gene-name with a semantically appropriate class\n(e.g. cst-pi-label) or remove the class if no styling is needed.',
        ],
        'status': 'Open',
    },
    {
        'id': 'K-11', 'severity': 'LOW', 'run': '1',
        'cols': [
            'fsn_hes_article03.html:1040\n(TOC close mechanism)',
            '<a href="#cst-toc-close">\n  <img id="cst-toc-close" …>\n(id placed on the img, not on a structural/focusable element)',
            'Move id="cst-toc-close" to the parent <a> or a containing <div>.\nId on an img used as a fragment target is fragile.',
        ],
        'status': 'Open',
    },
]

INFO_FINDINGS = [
    {
        'id': 'I-01', 'severity': 'INFO', 'run': '1',
        'cols': [
            'fsn_hes_article03.html\n(breadcrumb — client-managed per §4-4)',
            'Breadcrumb link points to /test.html (placeholder)',
            'Client-managed zone — not a deliverable defect.\nNote for client to update before production publish.',
        ],
        'status': 'N/A',
    },
]


# =============================================================================
# Documents
# =============================================================================

def create_overview():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, 'QA Overview — Project Cover Sheet')
    subtitle(doc, f'AZ HTML Production · {TICKET} · Chapter 1 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',  TICKET),
        ('Page / Component',  'fsn_hes_article03 — MediChannel AZ Japan ・ 特発性好酸球増加症候群診療の参照ガイド'),
        ('HTML file path',    'content/PhysicianServices/Japan/048-MediChannel/ja/jp/medical/product/fsn_hes-contents/fsn_hes_article03.html'),
        ('CSS files',         'etc/designs/code/…/fsn_hes_article01/base.css + page.css'),
        ('Figma file',        'ck8FNgtN3HLS0H4IR0Jx7W — AZ_04_HTML_v1.0'),
        ('Figma — Tab 1 desktop', 'node-id=8158-12921'),
        ('Figma — Tab 1 mobile SP', 'node-id=8158-13432'),
        ('Figma — Tab 2 desktop', 'node-id=8100-9345'),
        ('Figma — Tab 2 mobile SP', 'node-id=8100-9810'),
        ('QA date',           '2026-08-31'),
        ('QA reviewer',       'Claude AI (automated) — Run 1'),
        ('Scope exclusions',  'Header, breadcrumb, footer — client-managed per §4-4'),
    ])

    spacer(doc)
    section(doc, 'QA Chapter Summary')
    note(doc, 'Status after AI Run 1 (2026-08-31). Human review required to close findings.')

    hdrs = ['Chapter', 'Document', 'Status', 'Findings (open / total)', 'Signed off by', 'Date']
    ws   = [700, 2400, 900, 1800, 1872, 1400]
    rows_data = [
        ('Ch. 2', 'Design QA — WF vs Design Diff',       'N/A (no WF supplied)', '0 / 0', '', ''),
        ('Ch. 3', 'Content QA — Copy Accuracy',           'FAIL',                 '2 / 2', '', ''),
        ('Ch. 4', 'Coding QA — Typography & Compliance',  'FAIL',                 '11 / 11', '', ''),
        ('All',   'Overall QA result',                          'FAIL',                 '13 open', '', ''),
    ]
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(hdrs))
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)
    for i, row_vals in enumerate(rows_data):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, (val, w) in enumerate(zip(row_vals, ws)):
            clr = RED_C if val == 'FAIL' else None
            _data_cell(row.cells[j], val, w, bg, bold=(j == 0 or val == 'FAIL'), color=clr)
        _row_h(row, 480)

    spacer(doc)
    section(doc, 'Finding Summary by Severity')
    _t(doc, 'HIGH: 3 (K-01 empty title, K-02 duplicate main, K-03 missing h1) + 1 content (C-01 wrong hero alt)',
       size=9.5, color=RED_C, bold=True, before=2, after=2)
    _t(doc, 'HIGH total: 4 findings   —   MEDIUM: 5   —   LOW: 3   —   INFO: 1 (breadcrumb /test.html, client zone)',
       size=9.5, color=DARK_C, before=0, after=6)

    spacer(doc)
    section(doc, 'Revision History')
    note(doc, 'One row per QA run or significant change.')
    hdrs2 = ['Run', 'Date', 'Author', 'Summary of changes']
    ws2   = [700, 1600, 2200, PAGE_W - 700 - 1600 - 2200]
    tbl2  = doc.add_table(rows=4, cols=4)
    tbl2.style = 'Table Grid'
    _tbl_w(tbl2, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs2, ws2)):
        _hdr_cell(tbl2.rows[0].cells[i], h, w)
    _row_h(tbl2.rows[0], 420)
    run1_vals = ('1', '2026-08-31', 'Claude AI (automated)', 'Initial AI QA — 13 findings identified across content and coding')
    for i in range(1, 4):
        row = tbl2.rows[i]
        bg  = ALTROW if i % 2 == 1 else WHITE
        vals = run1_vals if i == 1 else ('', '', '', '')
        for j, (w, v) in enumerate(zip(ws2, vals)):
            _data_cell(row.cells[j], v, w, bg)
        _row_h(row, 460)

    return doc


def create_design_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, 'Design QA — WF vs Design Diff')
    subtitle(doc, f'AZ HTML Production · {TICKET} · Chapter 2 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',         TICKET),
        ('Page / Component',         'fsn_hes_article03 — 特発性好酸球増加症候群診療の参照ガイド'),
        ('Figma — Wireframe (WF)',     'Not supplied for this ticket'),
        ('Figma — Design Tab 1 PC',   'node-id=8158-12921'),
        ('Figma — Design Tab 1 SP',   'node-id=8158-13432'),
        ('Figma — Design Tab 2 PC',   'node-id=8100-9345'),
        ('Figma — Design Tab 2 SP',   'node-id=8100-9810'),
    ])

    spacer(doc)
    note(doc, 'No wireframe (WF) was supplied for this ticket. Design QA Chapter 2 is N/A for Run 1.')
    note(doc, 'If a WF is later provided, complete the checklist below against both WF and Design comps.')

    spacer(doc)
    section(doc, '2-1.  WF vs Design Diff')
    subsection(doc, 'Text & copy checks')
    check(doc, 'Heading text matches WF — no additions, deletions, or wording drift')
    check(doc, 'Body copy matches WF character-by-character')
    check(doc, 'Footnote and reference-number text matches WF')
    check(doc, 'Superscript / subscript / special characters (®, ™, °) match WF')
    check(doc, 'Full-width vs half-width characters correct (Japanese pages)')
    check(doc, 'No content present in WF is missing from Design')
    check(doc, 'No content added to Design that was not in WF')

    subsection(doc, 'Image & chart checks')
    check(doc, 'Graph / chart data labels and values match WF')
    check(doc, 'Chart footnote numbers match WF')
    check(doc, 'Image subject and composition match WF intent')

    spacer(doc)
    section(doc, '2-2.  Design Guideline Compliance')
    check(doc, 'Brand colors match guideline spec (Primary: #ae2573, Secondary: #00b398)')
    check(doc, 'Typography: font family and weight match guideline')
    check(doc, 'Logo: clear-space rules observed')
    check(doc, 'No custom or off-palette colors introduced')

    spacer(doc)
    section(doc, 'Session History')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings')
    note(doc, 'Run 1 (AI): No WF supplied — no findings raised for Chapter 2. Status: N/A.')
    note(doc, 'Status: Open · Fixed (Run N) · Accepted · N/A')
    findings_table_populated(
        doc,
        headers=['Location', 'WF side', 'Design side'],
        col_widths=[1800, 2500, 2500],
        findings=[],
    )
    return doc


def create_content_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, 'Content QA — Copy Accuracy (Design vs Code)')
    subtitle(doc, f'AZ HTML Production · {TICKET} · Chapter 3 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',  TICKET),
        ('Page / Component',  'fsn_hes_article03 — 特発性好酸球増加症候群診療の参照ガイド'),
        ('HTML file path',    'content/PhysicianServices/Japan/048-MediChannel/ja/jp/medical/product/fsn_hes-contents/fsn_hes_article03.html'),
        ('Figma — Design Tab 1 PC',  'node-id=8158-12921'),
        ('Figma — Design Tab 1 SP',  'node-id=8158-13432'),
        ('Figma — Design Tab 2 PC',  'node-id=8100-9345'),
        ('Figma — Design Tab 2 SP',  'node-id=8100-9810'),
    ])

    spacer(doc)
    _t(doc, 'Scope: text and copy only — no layout or styling. Compare each visible string in the coded HTML against the Figma design comp character-by-character.',
       size=9.5, italic=True, color=GREY_C, before=4, after=4, indent=0.15)

    section(doc, '3-1.  Design vs Code Copy Match')
    subsection(doc, 'Headings')
    check(doc, 'H1 heading text matches Design exactly (including punctuation and spacing)')
    check(doc, 'H2 heading text matches Design')
    check(doc, 'H3 / H4 heading text matches Design')
    check(doc, 'CQ tab / section tab labels match Design')

    subsection(doc, 'Body copy')
    check(doc, 'Body paragraph text matches Design character-by-character')
    check(doc, 'List items match Design (order and wording)')
    check(doc, 'Button and link labels match Design')

    subsection(doc, 'Footnotes & references')
    check(doc, 'Footnote superscript numbers match Design')
    check(doc, 'Footnote body text matches Design')
    check(doc, 'Reference numbers and text match Design')

    subsection(doc, 'Special characters')
    check(doc, '® / ™ / © marks present and correctly placed')
    check(doc, 'Superscripts coded as <sup>, not as Unicode characters')
    check(doc, 'Full-width numerals used where Design specifies (Japanese pages)')
    check(doc, 'Japanese en-dashes (—) and ellipsis match Design')

    subsection(doc, 'Completeness')
    check(doc, 'No content present in Design is missing from HTML')
    check(doc, 'No extra content in HTML that is not in Design')

    spacer(doc)
    section(doc, '3-2.  Placeholder Detection')
    note(doc, 'Any of the below found = BLOCKER. Must be resolved before delivery.')
    check(doc, 'No breadcrumb link pointing to /test.html  → INFO only per §4-4 (client-managed)')
    check(doc, 'No approval code placeholder (JP-○○○○) in .cst-page-info')
    check(doc, 'No “Lorem”, “TODO”, “PLACEHOLDER”, “SAMPLE”, or “DUMMY” text')

    spacer(doc)
    section(doc, 'Session History')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings  —  Run 1 (AI, 2026-08-31)')
    note(doc, 'Sev column: HIGH · MEDIUM · LOW · INFO   |   Status: Open · Fixed (Run N) · Accepted · N/A')
    findings_table_populated(
        doc,
        headers=['Location', 'Design side (Figma)', 'HTML side (current code)'],
        col_widths=[1800, 2500, 2500],
        findings=CONTENT_FINDINGS + INFO_FINDINGS,
    )
    spacer(doc)
    note(doc, 'Total: 2 open findings (1 HIGH, 1 MEDIUM) + 1 INFO (client-managed breadcrumb placeholder).')
    return doc


def create_coding_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, 'Coding QA — Typography Metrics & Guideline Compliance')
    subtitle(doc, f'AZ HTML Production · {TICKET} · Chapter 4 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',  TICKET),
        ('Page / Component',  'fsn_hes_article03 — 特発性好酸球増加症候群診療の参照ガイド'),
        ('HTML file path',    'content/PhysicianServices/Japan/048-MediChannel/ja/jp/medical/product/fsn_hes-contents/fsn_hes_article03.html'),
        ('CSS files',         'etc/designs/code/…/fsn_hes_article01/base.css  +  page.css'),
        ('Figma — Design Tab 1 PC',  'node-id=8158-12921'),
        ('Figma — Design Tab 1 SP',  'node-id=8158-13432'),
    ])

    spacer(doc)
    _t(doc, 'Scope: technical and visual-metric compliance only. Copy and text accuracy is covered in Content QA (Chapter 3).',
       size=9.5, italic=True, color=GREY_C, before=4, after=4, indent=0.15)

    section(doc, '4-1.  Typography Metrics')
    note(doc, 'Browser values from DevTools Computed tab with live preview server. Fill in after human browser check.')
    hdrs = ['Element', 'Figma (PC)', 'Browser (PC)', 'Figma (SP)', 'Browser (SP)', 'Pass / Fail']
    ws   = [1372, 1500, 1600, 1500, 1600, 1500]
    elms = ['H1', 'H2', 'H3', 'H4', 'Body text', 'Footnote', 'Ref / caption', 'Tab label']
    tbl  = doc.add_table(rows=1 + len(elms), cols=6)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)
    for i, el in enumerate(elms):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, w in enumerate(ws):
            _data_cell(row.cells[j], el if j == 0 else '', w, bg, bold=(j == 0))
        _row_h(row, 460)

    spacer(doc)
    section(doc, '4-2.  Coding Guideline Compliance')
    note(doc, 'Reference: guidelines/base/xhtml-coding-rules.md + guidelines/base/medichannel-delivery-standards.md')
    note(doc, 'Do NOT use html-coding-rules.md for MediChannel — XHTML 1.0 Strict rules apply.')

    subsection(doc, 'XHTML document structure')
    check(doc, '<?xml version="1.0" encoding="UTF-8"?> present at line 1')
    check(doc, 'DOCTYPE is XHTML 1.0 Strict')
    check(doc, 'Content-Type meta uses http-equiv, charset=UTF-8, placed first in <head>')
    check(doc, 'html[lang] set to correct page language')
    check(doc, '✖  <title> is unique, descriptive, 45–65 characters  →  FAIL (K-01: empty)')
    check(doc, 'All tags properly closed (void elements: <br />, <img />, <input />)')
    check(doc, 'All attribute values quoted with double quotes')

    subsection(doc, 'Semantics & accessibility')
    check(doc, '✖  Exactly one <main> / role="main" landmark  →  FAIL (K-02: two role="main")')
    check(doc, '✖  Exactly one h1 on the page  →  FAIL (K-03: no h1 at all)')
    check(doc, 'Heading hierarchy correct — no skipped levels (h1 → h2 → h3)')
    check(doc, '✖  All images have meaningful alt text  →  FAIL (C-01: hero alt wrong, see Content QA)')
    check(doc, 'All form inputs have associated <label> elements')
    check(doc, '✖  Navigation landmarks have accessible names  →  FAIL (K-09: secondary TOC nav missing role)')
    check(doc, 'Visible focus states present on all keyboard-focusable elements')

    subsection(doc, 'Image loading')
    check(doc, '✖  LCP/hero images use loading="eager" fetchpriority="high"  →  FAIL (K-04: lazy on hero)')

    subsection(doc, 'CSS hygiene')
    check(doc, 'No commented-out CSS blocks')
    check(doc, 'No hardcoded hex colors outside CSS custom properties (:root)')
    check(doc, 'No !important declarations')
    check(doc, '✖  No structural class reuse causing incorrect side-effects  →  FAIL (K-05: cst-hero-wrapper)')
    check(doc, '✖  scroll-margin-top set on sticky-header anchor targets  →  FAIL (K-06: missing)')
    check(doc, "font-family stack includes 'Meiryo' or 'Hiragino Kaku Gothic ProN' (Japanese pages)")

    subsection(doc, 'Semantic element use')
    check(doc, '✖  Links must use <a href> not styled <p> or <div>  →  FAIL (K-07: Tab 2 footer "link")')
    check(doc, '✖  Approval notice element type consistent across tabs  →  FAIL (K-08: strong vs h4)')
    check(doc, '✖  Class names match semantic purpose of element  →  FAIL (K-10: cst-gene-name misused)')
    check(doc, '✖  Fragment-target ids on structural elements, not img  →  FAIL (K-11: id on img)')

    subsection(doc, 'MediChannel delivery')
    check(doc, 'Editable area confined to .cst-page (no edits outside)')
    check(doc, 'Header / footer / nav untouched (client-managed per §4-4)')
    check(doc, 'jQuery version 1.8.3 — no other version loaded')
    check(doc, 'Total file size (HTML + CSS) ≤ 800 KB')
    check(doc, 'Approved AZ MediChannel template used as base')

    spacer(doc)
    section(doc, 'Session History')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings  —  Run 1 (AI, 2026-08-31)')
    note(doc, 'Sev column: HIGH · MEDIUM · LOW · INFO   |   Status: Open · Fixed (Run N) · Accepted · N/A')
    findings_table_populated(
        doc,
        headers=['File + Line', 'Current code', 'Corrected code'],
        col_widths=[1800, 2500, 2500],
        findings=CODING_FINDINGS,
    )
    spacer(doc)
    note(doc, 'Total: 11 open findings (3 HIGH, 5 MEDIUM, 3 LOW).')
    return doc


# =============================================================================
# Main
# =============================================================================
if __name__ == '__main__':
    os.makedirs(OUT_DIR, exist_ok=True)
    files = [
        (f'{TICKET}_overview-qa.docx',  create_overview()),
        (f'{TICKET}_design-qa.docx',    create_design_qa()),
        (f'{TICKET}_content-qa.docx',   create_content_qa()),
        (f'{TICKET}_coding-qa.docx',    create_coding_qa()),
    ]
    for name, doc in files:
        path = os.path.join(OUT_DIR, name)
        doc.save(path)
        print(f'Written: {path}')
