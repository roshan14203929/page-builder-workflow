"""
Create QA template DOCX files for one ticket:
  qa-reports/<TICKET>/
    <TICKET>_design-qa.docx
    <TICKET>_content-qa.docx
    <TICKET>_coding-qa.docx

Usage:
  python create_qa_docs.py AGENCY15-316
  python create_qa_docs.py            # uses TICKET-ID placeholder
"""
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

QA_ROOT = str(Path(__file__).parent.parent / 'qa-reports')
TICKET  = sys.argv[1] if len(sys.argv) > 1 else 'TICKET-ID'

# --- Colors ---
NAVY   = '1F4E79'
BLUE   = '2E86AB'
LGREY  = 'D5D8DC'
ALTROW = 'EBF5FB'
WHITE  = 'FFFFFF'
NAVY_C  = RGBColor(0x1F, 0x4E, 0x79)
BLUE_C  = RGBColor(0x2E, 0x86, 0xAB)
WHITE_C = RGBColor(0xFF, 0xFF, 0xFF)
DARK_C  = RGBColor(0x1A, 0x1A, 0x1A)
GREY_C  = RGBColor(0x66, 0x66, 0x66)
LGREY_C = RGBColor(0x88, 0x88, 0x88)

# --- XML helpers ---
def _cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _cell_w(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:w'), str(dxa))
    tcW.set(qn('w:type'), 'dxa')

def _tbl_w(table, dxa):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(dxa))
    tblW.set(qn('w:type'), 'dxa')

def _row_h(row, val=380):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight')
    h.set(qn('w:val'), str(val))
    h.set(qn('w:hRule'), 'atLeast')
    trPr.append(h)

def _para_border_bottom(para, color=NAVY):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

# --- Text helpers ---
def _t(doc, text, size=9.5, bold=False, italic=False, color=None,
       before=0, after=2, indent=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    if align:  p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def title(doc, text):
    return _t(doc, text, size=20, bold=True, color=NAVY_C, before=0, after=4)

def subtitle(doc, text):
    return _t(doc, text, size=10, italic=True, color=GREY_C, before=0, after=10)

def section(doc, text):
    p = _t(doc, text, size=12, bold=True, color=NAVY_C, before=12, after=3)
    _para_border_bottom(p)
    return p

def subsection(doc, text):
    return _t(doc, text, size=10, bold=True, color=BLUE_C, before=6, after=2)

def check(doc, text):
    return _t(doc, '\u2610  ' + text, size=9.5, color=DARK_C, before=1, after=1, indent=0.2)

def note(doc, text):
    return _t(doc, text, size=8.5, italic=True, color=LGREY_C, before=2, after=2, indent=0.2)

def spacer(doc):
    _t(doc, '', size=8, before=0, after=4)

# --- Table helpers ---
PAGE_W = 9072   # DXA: A4 minus 1-inch margins each side

def _hdr_cell(cell, text, w):
    _cell_w(cell, w); _cell_bg(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE_C

def _data_cell(cell, text, w, bg=WHITE, bold=False):
    _cell_w(cell, w); _cell_bg(cell, bg)
    if text:
        r = cell.paragraphs[0].add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = GREY_C
        if bold: r.bold = True

def info_table(doc, rows):
    W = [2000, PAGE_W - 2000]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (label, val) in enumerate(rows):
        r = tbl.rows[i]
        _cell_w(r.cells[0], W[0]); _cell_bg(r.cells[0], LGREY)
        _cell_w(r.cells[1], W[1]); _cell_bg(r.cells[1], WHITE)
        rr = r.cells[0].paragraphs[0].add_run(label)
        rr.bold = True; rr.font.size = Pt(9)
        if val:
            rr2 = r.cells[1].paragraphs[0].add_run(val)
            rr2.font.size = Pt(9)
        _row_h(r, 340)

def session_log(doc, n=3):
    """Session history table: Run | Date | Reviewed by | Notes."""
    W = [700, 2000, 2500, PAGE_W - 700 - 2000 - 2500]
    hdrs = ['Run', 'Date', 'Reviewed by', 'Notes']
    tbl = doc.add_table(rows=1 + n, cols=4)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, W)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 400)
    for i in range(n):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, w in enumerate(W):
            _data_cell(row.cells[j], str(i + 1) if j == 0 else '', w, bg)
        _row_h(row, 420)

def findings_table(doc, headers, col_widths, id_prefix, n=15):
    """
    Findings table with a 'Run' column so reviewers mark which QA session
    introduced each finding.  Columns: ID | Run | <phase headers> | Status
    """
    all_hdrs   = ['ID', 'Run'] + headers + ['Status']
    run_w      = 580
    status_w   = 900
    inner_total = PAGE_W - 700 - run_w - status_w
    # Distribute inner_total across phase-specific columns proportionally
    ratio       = inner_total / sum(col_widths)
    inner_ws    = [max(800, int(w * ratio)) for w in col_widths]
    # Correct rounding drift on last col
    inner_ws[-1] += inner_total - sum(inner_ws)
    all_ws = [700, run_w] + inner_ws + [status_w]

    tbl = doc.add_table(rows=1 + n, cols=len(all_hdrs))
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)

    for i, (h, w) in enumerate(zip(all_hdrs, all_ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)

    for i in range(n):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        id_text = f'{id_prefix}-{i+1:02d}'
        for j, w in enumerate(all_ws):
            text = id_text if j == 0 else ''
            _data_cell(row.cells[j], text, w, bg)
        _row_h(row, 500)

def typo_table(doc):
    hdrs = ['Element', 'Figma (PC)', 'Browser (PC)', 'Figma (SP)', 'Browser (SP)', 'Pass / Fail']
    ws   = [1372, 1500, 1600, 1500, 1600, 1500]
    elms = ['H1', 'H2', 'H3', 'H4', 'Body text', 'Footnote', 'Ref / caption', 'Tab label']
    tbl  = doc.add_table(rows=1 + len(elms), cols=6)
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)
    for i, el in enumerate(elms):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, w in enumerate(ws):
            _data_cell(row.cells[j], el if j == 0 else '', w, bg, bold=(j == 0))
        _row_h(row, 460)

# =============================================================================
# 1. DESIGN QA
# =============================================================================
def create_design_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, f'Design QA \u2014 WF vs Design Diff')
    subtitle(doc, f'AZ HTML Production \u00b7 {TICKET} \u00b7 Chapter 2 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',          TICKET),
        ('Page / Component',          ''),
        ('Figma \u2014 Wireframe (WF)',      ''),
        ('Figma \u2014 Design PC',           ''),
        ('Figma \u2014 Design SP',           ''),
    ])

    spacer(doc)
    section(doc, '2-1.  WF vs Design Diff')
    subsection(doc, 'Text & copy checks')
    check(doc, 'Heading text matches WF \u2014 no additions, deletions, or wording drift')
    check(doc, 'Body copy matches WF character-by-character')
    check(doc, 'Footnote and reference-number text matches WF')
    check(doc, 'Superscript / subscript / special characters (\u00ae, \u2122, \u00b0) match WF')
    check(doc, 'Full-width vs half-width characters correct (Japanese pages)')
    check(doc, 'No content present in WF is missing from Design')
    check(doc, 'No content added to Design that was not in WF')

    subsection(doc, 'Image & chart checks')
    check(doc, 'Graph / chart data labels and values match WF')
    check(doc, 'Chart footnote numbers match WF')
    check(doc, 'Image subject and composition match WF intent')

    spacer(doc)
    section(doc, '2-2.  Design Guideline Compliance')
    note(doc, 'Reference: brand guideline page in Figma. Ask the project lead for the current link.')
    check(doc, 'Brand colors match guideline spec (Primary and Secondary palette)')
    check(doc, 'Typography: font family and weight match guideline')
    check(doc, 'Logo: clear-space rules observed')
    check(doc, 'Logo: correct naming variant used')
    check(doc, 'No custom or off-palette colors introduced')

    spacer(doc)
    section(doc, 'Session History')
    note(doc, 'Fill in one row per QA run. Run number matches the Run column in the findings table below.')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings')
    note(doc, 'Run column: enter the run number when the finding was first raised (e.g. 1, 2, 3).')
    note(doc, 'Status: Open \u00b7 Fixed (Run N) \u00b7 Accepted \u00b7 N/A')
    findings_table(
        doc,
        headers=['Location', 'WF side', 'Design side'],
        col_widths=[1800, 2500, 2500],
        id_prefix='D',
    )
    spacer(doc)
    note(doc, 'Severity: BLOCKER \u00b7 HIGH \u00b7 MEDIUM \u00b7 LOW \u00b7 INFO')
    return doc

# =============================================================================
# 2. CONTENT QA
# =============================================================================
def create_content_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, f'Content QA \u2014 Copy Accuracy (Design vs Code)')
    subtitle(doc, f'AZ HTML Production \u00b7 {TICKET} \u00b7 Chapter 3 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',     TICKET),
        ('Page / Component',     ''),
        ('HTML file path',       ''),
        ('Figma \u2014 Design PC',      ''),
        ('Figma \u2014 Design SP',      ''),
    ])

    spacer(doc)
    _t(doc, 'Scope: text and copy only \u2014 no layout or styling. Compare each visible string in the coded HTML against the Figma design comp character-by-character.',
       size=9.5, italic=True, color=GREY_C, before=4, after=4, indent=0.15)

    section(doc, '3-1.  Design vs Code Copy Match')
    subsection(doc, 'Headings')
    check(doc, 'H1 heading text matches Design exactly (including punctuation and spacing)')
    check(doc, 'H2 heading text matches Design')
    check(doc, 'H3 / H4 heading text matches Design')
    check(doc, 'CQ tab / section tab labels match Design')

    subsection(doc, 'Body copy')
    check(doc, 'Body paragraph text matches Design character-by-character')
    check(doc, 'List items match Design (order and wording)')
    check(doc, 'Button and link labels match Design')
    check(doc, 'Form field labels and placeholder text match Design')

    subsection(doc, 'Footnotes & references')
    check(doc, 'Footnote superscript numbers match Design')
    check(doc, 'Footnote body text matches Design')
    check(doc, 'Reference numbers and text match Design')

    subsection(doc, 'Special characters')
    check(doc, '\u00ae / \u2122 / \u00a9 marks present and correctly placed')
    check(doc, 'Superscripts coded as <sup>, not as Unicode characters')
    check(doc, 'Subscripts coded as <sub>')
    check(doc, 'Full-width numerals used where Design specifies them (Japanese pages)')
    check(doc, 'Half-width numerals used where Design specifies them')
    check(doc, 'Japanese en-dashes (\u2014) and ellipsis (\u2026\u2026) match Design')

    subsection(doc, 'Completeness')
    check(doc, 'No content present in Design is missing from HTML')
    check(doc, 'No extra content in HTML that is not in Design')

    spacer(doc)
    section(doc, '3-2.  Placeholder Detection')
    note(doc, 'Any of the below found = BLOCKER. Must be resolved before delivery.')
    check(doc, 'No breadcrumb link pointing to /test.html')
    check(doc, 'No breadcrumb visible text containing \u201c test\u201d')
    check(doc, 'No approval code placeholder (JP-\u25cb\u25cb\u25cb\u25cb) in .cst-page-info')
    check(doc, 'No approval code placeholder in footer note element (.cst-footer-note__code)')
    check(doc, 'No \u201cLorem\u201d, \u201cTODO\u201d, \u201cPLACEHOLDER\u201d, \u201cSAMPLE\u201d, or \u201cDUMMY\u201d text')
    check(doc, 'No Japanese placeholder text (\u3053\u3053\u306b\u5165\u308b / \u30c6\u30b9\u30c8 / \u30b5\u30f3\u30d7\u30eb) visible in DOM')

    spacer(doc)
    section(doc, 'Session History')
    note(doc, 'Fill in one row per QA run. Run number matches the Run column in the findings table below.')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings')
    note(doc, 'Run column: enter the run number when the finding was first raised (e.g. 1, 2, 3).')
    note(doc, 'Status: Open \u00b7 Fixed (Run N) \u00b7 Accepted \u00b7 N/A')
    findings_table(
        doc,
        headers=['Location', 'Design side', 'HTML side'],
        col_widths=[1800, 2500, 2500],
        id_prefix='C',
    )
    spacer(doc)
    note(doc, 'Severity: BLOCKER \u00b7 HIGH \u00b7 MEDIUM \u00b7 LOW \u00b7 INFO')
    return doc

# =============================================================================
# 3. CODING QA
# =============================================================================
def create_coding_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, f'Coding QA \u2014 Typography Metrics & Guideline Compliance')
    subtitle(doc, f'AZ HTML Production \u00b7 {TICKET} \u00b7 Chapter 4 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',     TICKET),
        ('Page / Component',     ''),
        ('HTML file path',       ''),
        ('CSS files',            ''),
        ('Figma \u2014 Design PC',      ''),
        ('Figma \u2014 Design SP',      ''),
    ])

    spacer(doc)
    _t(doc, 'Scope: technical and visual-metric compliance only. Copy and text accuracy is covered in Content QA (Chapter 3).',
       size=9.5, italic=True, color=GREY_C, before=4, after=4, indent=0.15)

    section(doc, '4-1.  Typography Metrics')
    note(doc, 'Reference: Figma design comp (Dev mode). Browser values from DevTools Computed tab with the preview server running.')
    subsection(doc, 'Measurement table  (font-size / line-height in px)')
    typo_table(doc)

    spacer(doc)
    section(doc, '4-2.  Coding Guideline Compliance')
    note(doc, 'Reference: guidelines/base/xhtml-coding-rules.md + guidelines/base/medichannel-delivery-standards.md')
    note(doc, 'Do NOT use html-coding-rules.md or /html-css-review for MediChannel \u2014 those are for 3rd Party (HTML5) only.')

    subsection(doc, 'XHTML document structure')
    check(doc, '<?xml version="1.0" encoding="UTF-8"?> present at line 1')
    check(doc, 'DOCTYPE is XHTML 1.0 Strict')
    check(doc, 'Content-Type meta uses http-equiv (not name), charset=UTF-8, placed first in <head>')
    check(doc, 'html[lang] set to correct page language')
    check(doc, '<title> is unique, descriptive, 45\u201365 characters')
    check(doc, 'All tags properly closed (void elements: <br />, <img />, <input />, etc.)')
    check(doc, 'All attribute values quoted with double quotes')
    check(doc, 'id and class values are lowercase-hyphenated (no camelCase)')

    subsection(doc, 'XHTML scripting')
    check(doc, 'No inline <script> blocks without CDATA wrapping (// <![CDATA[ \u2026 // ]]>)')
    check(doc, 'No inline event handlers (onclick, onload) \u2014 event listeners in JS only')

    subsection(doc, 'Semantics & accessibility')
    check(doc, 'Exactly one <main> element (or role="main") \u2014 no duplicates')
    check(doc, 'Heading hierarchy correct \u2014 no skipped levels (h1 \u2192 h2 \u2192 h3)')
    check(doc, 'Exactly one h1 on the page')
    check(doc, 'All images have meaningful alt text (decorative: alt="")')
    check(doc, 'All form inputs have associated <label> elements')
    check(doc, 'Interactive controls have accessible names')
    check(doc, 'Visible focus states present on all keyboard-focusable elements')

    subsection(doc, 'CSS hygiene')
    check(doc, 'No commented-out CSS blocks')
    check(doc, 'No hardcoded hex colors outside CSS custom properties (:root variables)')
    check(doc, 'No !important declarations')
    check(doc, 'No duplicate selectors that cancel each other out')
    check(doc, "font-family stack includes 'Hiragino Kaku Gothic ProN' (Japanese pages)")

    subsection(doc, 'MediChannel delivery')
    check(doc, 'Editable area confined to .cst-page (no edits outside)')
    check(doc, 'Header / footer / nav untouched')
    check(doc, 'jQuery version 1.8.3 \u2014 no other version loaded')
    check(doc, 'Total file size (HTML + CSS) \u2264 800 KB')
    check(doc, 'Template DOCTYPE and charset preserved')
    check(doc, 'Approved AZ delivery template used as base')

    spacer(doc)
    section(doc, 'Session History')
    note(doc, 'Fill in one row per QA run. Run number matches the Run column in the findings table below.')
    session_log(doc)

    spacer(doc)
    section(doc, 'Findings')
    note(doc, 'Run column: enter the run number when the finding was first raised (e.g. 1, 2, 3).')
    note(doc, 'Status: Open \u00b7 Fixed (Run N) \u00b7 Accepted \u00b7 N/A')
    findings_table(
        doc,
        headers=['File + Line', 'Current code', 'Corrected code'],
        col_widths=[1800, 2500, 2500],
        id_prefix='K',
    )
    spacer(doc)
    note(doc, 'Severity: BLOCKER \u00b7 HIGH \u00b7 MEDIUM \u00b7 LOW \u00b7 INFO')
    return doc

# =============================================================================
# 0. OVERVIEW (Chapter 1) — project cover sheet + QA sign-off
# =============================================================================
def create_overview_qa():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    title(doc, f'QA Overview — Project Cover Sheet')
    subtitle(doc, f'AZ HTML Production · {TICKET} · Chapter 1 of the QA Guide')

    section(doc, 'Project Information')
    info_table(doc, [
        ('Ticket / Project',     TICKET),
        ('Page / Component',     ''),
        ('HTML file path',       ''),
        ('Preview URL',          ''),
        ('Figma file',           ''),
        ('Figma — WF PC',         ''),
        ('Figma — WF SP',         ''),
        ('Figma — Design PC',     ''),
        ('Figma — Design SP',     ''),
        ('Delivery date',        ''),
        ('Developer',            ''),
        ('QA reviewer',          ''),
    ])

    spacer(doc)
    section(doc, 'QA Chapter Summary')
    note(doc, 'Update after each QA run. Status: — (not started) · In Progress · Pass · Fail · N/A')

    hdrs = ['Chapter', 'Document', 'Status', 'Findings (open / total)', 'Signed off by', 'Date']
    ws   = [700, 2400, 900, 1800, 1872, 1400]
    rows_data = [
        ('Ch. 2', 'Design QA — WF vs Design Diff',            '', '', '', ''),
        ('Ch. 3', 'Content QA — Copy Accuracy',               '', '', '', ''),
        ('Ch. 4', 'Coding QA — Typography & Compliance',      '', '', '', ''),
        ('All',   'Overall QA result',                             '', '', '', ''),
    ]
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(hdrs))
    tbl.style = 'Table Grid'
    _tbl_w(tbl, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs, ws)):
        _hdr_cell(tbl.rows[0].cells[i], h, w)
    _row_h(tbl.rows[0], 420)
    for i, row_vals in enumerate(rows_data):
        row = tbl.rows[i + 1]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, (val, w) in enumerate(zip(row_vals, ws)):
            _data_cell(row.cells[j], val, w, bg, bold=(j == 0))
        _row_h(row, 480)

    spacer(doc)
    section(doc, 'Revision History')
    note(doc, 'One row per QA run or significant change.')
    hdrs2 = ['Run', 'Date', 'Author', 'Summary of changes']
    ws2   = [700, 1600, 2200, PAGE_W - 700 - 1600 - 2200]
    tbl2  = doc.add_table(rows=4, cols=4)
    tbl2.style = 'Table Grid'
    _tbl_w(tbl2, PAGE_W)
    for i, (h, w) in enumerate(zip(hdrs2, ws2)):
        _hdr_cell(tbl2.rows[0].cells[i], h, w)
    _row_h(tbl2.rows[0], 420)
    for i in range(1, 4):
        row = tbl2.rows[i]
        bg  = ALTROW if i % 2 == 1 else WHITE
        for j, w in enumerate(ws2):
            _data_cell(row.cells[j], str(i) if j == 0 else '', w, bg)
        _row_h(row, 460)

    spacer(doc)
    section(doc, 'Notes')
    _t(doc, '', size=9.5, before=0, after=60)

    return doc


# =============================================================================
# Main
# =============================================================================
if __name__ == '__main__':
    out_dir = os.path.join(QA_ROOT, TICKET)
    os.makedirs(out_dir, exist_ok=True)

    files = [
        (f'{TICKET}_overview-qa.docx', create_overview_qa()),
        (f'{TICKET}_design-qa.docx',   create_design_qa()),
        (f'{TICKET}_content-qa.docx',  create_content_qa()),
        (f'{TICKET}_coding-qa.docx',   create_coding_qa()),
    ]
    for name, doc in files:
        path = os.path.join(out_dir, name)
        doc.save(path)
        print(f'Written: {path}')
